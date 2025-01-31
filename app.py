from flask import Flask, request, jsonify, render_template, send_file
import os
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
import json
import fitz 
from docx import Document
from collections import defaultdict

app = Flask(__name__)

# Google Drive API Scopes
SCOPES = ['https://www.googleapis.com/auth/drive.file']
CREDENTIALS_FILE = 'credentials.json'
TOKEN_FILE = 'token.json'
LOCAL_DOCS_DIR = 'local_docs'

# Use defaultdict to avoid KeyError when accessing non-existent files
last_added_text = defaultdict(list)  # Change to list to maintain history
MAX_UNDO_HISTORY = 10  # Keep last 10 changes

os.makedirs(LOCAL_DOCS_DIR, exist_ok=True)

def authenticate_google_drive():
    """Authenticate with Google Drive API and return the service object."""
    creds = None
    # Load existing tokens if available
    if os.path.exists(TOKEN_FILE):
        with open(TOKEN_FILE, 'r') as token:
            creds = Credentials.from_authorized_user_info(json.load(token), SCOPES)
    
    # If no valid credentials, prompt the user to log in
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(CREDENTIALS_FILE, SCOPES)
            creds = flow.run_local_server(port=0)
        
        # Save the credentials for future use
        with open(TOKEN_FILE, 'w') as token:
            token.write(creds.to_json())
    
    # Build the Google Drive service
    service = build('drive', 'v3', credentials=creds)
    return service

@app.route('/extract-pdf-text', methods=['POST'])
def extract_pdf_text():
    """Endpoint to extract text from a PDF file."""
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    # Save the uploaded file temporarily
    temp_pdf_path = os.path.join(LOCAL_DOCS_DIR, "temp.pdf")
    file.save(temp_pdf_path)

    # Extract text from the PDF
    try:
        text = ""
        with fitz.open(temp_pdf_path) as doc:
            for page in doc:
                text += page.get_text()
        
        # Delete the temporary file
        os.remove(temp_pdf_path)

        return jsonify({'text': text}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def create_google_doc(service, text, title="Saved PDF Text"):
    """Create a Google Doc with the given text and return the file ID."""
    # Create a file metadata
    
    file_metadata = {
        'name': title,
        'mimeType': 'application/vnd.google-apps.document'
    }
    
    # Create the file on Google Drive
    
    file = service.files().create(body=file_metadata, media_body=None).execute()
    
    # Update the file with the provided text
    service.files().update(fileId=file['id'], media_body=text).execute()
    
    return file['id']

def create_local_docx(text, file_name="Saved PDF Text", save_location=None):
    """Create or update a local .docx file with the given text and return the file path."""
    if not file_name.lower().endswith(".docx"):
        file_name += ".docx"

    save_dir = save_location if save_location and os.path.isdir(save_location) else LOCAL_DOCS_DIR
    os.makedirs(save_dir, exist_ok=True)
    file_path = os.path.join(save_dir, file_name)

    try:
        if os.path.exists(file_path):
            doc = Document(file_path)
            # Add a blank line before new content if the document isn't empty
            # if len(doc.paragraphs) > 0 and doc.paragraphs[-1].text.strip() != "":
            #     doc.add_paragraph("")  # Add blank line for spacing
        else:
            doc = Document()
        
        # Add the new text
        doc.add_paragraph(text)
        
        # Store the text in undo history
        if len(last_added_text[file_path]) >= MAX_UNDO_HISTORY:
            last_added_text[file_path].pop(0)  # Remove oldest entry if limit reached
        last_added_text[file_path].append(text)
        
        # Save the document
        doc.save(file_path)
        
        return {"file_path": file_path, "message": "Text saved successfully!"}
    except Exception as e:
        return {"error": f"Failed to save document: {str(e)}"}
    
# @app.route("/")
# def hello():
#     return "Hello World!"

@app.route("/")
def home():
    return render_template("index.html")  # Serve the frontend

@app.route('/list-files')
def list_files():
    """Endpoint to list all .docx files in the local_docs directory."""
    files = [f for f in os.listdir(LOCAL_DOCS_DIR) if f.endswith('.docx')]
    return jsonify(files)

@app.route('/get-file')
def get_file():
    """Endpoint to get the content of a specific .docx file."""
    file_name = request.args.get('name')
    print("line142", LOCAL_DOCS_DIR, file_name)
    file_path = os.path.join(LOCAL_DOCS_DIR, file_name)
    
    if not os.path.exists(file_path):
        return jsonify({'error': 'File not found'}), 404
    
    doc = Document(file_path)
    content = '\n'.join([para.text for para in doc.paragraphs])
    return jsonify({'content': content})

@app.route('/view-file')
def view_file():
    """Endpoint to view file content with proper formatting"""
    file_name = request.args.get('name')
    file_path = os.path.join(LOCAL_DOCS_DIR, file_name)
    
    if not os.path.exists(file_path):
        return jsonify({'error': 'File not found'}), 404

    try:
        doc = Document(file_path)
        # Join paragraphs with newlines to preserve formatting
        content = '\n'.join(para.text for para in doc.paragraphs)
        return content
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/create-file', methods=['POST'])
def create_file():
    """Endpoint to create a new empty .docx file."""
    data = request.json
    file_name = data.get('fileName', 'Untitled.docx')
    
    if not file_name.lower().endswith(".docx"):
        file_name += ".docx"
    
    file_path = os.path.join(LOCAL_DOCS_DIR, file_name)
    
    try:
        if os.path.exists(file_path):
            return jsonify({'success': False, 'message': 'File already exists'}), 400
        
        doc = Document()
        doc.save(file_path)
        return jsonify({'success': True, 'message': 'File created successfully'}), 200
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/undo-last-text', methods=['POST'])
def undo_last_text():
    """Endpoint to remove the last added text from a file."""
    try:
        data = request.json
        file_name = data.get('file_name')

        if not file_name:
            return jsonify({'error': 'File name not provided'}), 400

        if not file_name.lower().endswith(".docx"):
            file_name += ".docx"
        print("local_docs",LOCAL_DOCS_DIR, file_name )
        file_path = os.path.join(LOCAL_DOCS_DIR, file_name)

        if not os.path.exists(file_path):
            return jsonify({'error': 'File not found'}), 404

        # Check if there's any undo history for this file
        if not last_added_text[file_path]:
            return jsonify({'error': 'No text to undo'}), 400

        # Load the document
        doc = Document(file_path)
        
        if not doc.paragraphs:
            return jsonify({'error': 'Document is already empty'}), 400

        # Get the last added text from our history
        last_text = last_added_text[file_path][-1]
        
        # Find and remove the last paragraph that matches our stored text
        found = False
        for i in range(len(doc.paragraphs) - 1, -1, -1):
            if doc.paragraphs[i].text == last_text:
                p = doc.paragraphs[i]._element
                p.getparent().remove(p)
                found = True
                break

        if not found:
            return jsonify({'error': 'Last added text not found in file'}), 400

        # Save the document and update our history
        doc.save(file_path)
        last_added_text[file_path].pop()  # Remove the last entry from history

        return jsonify({
            'message': 'Last added text removed successfully',
            'remaining_undos': len(last_added_text[file_path])
        }), 200

    except Exception as e:
        return jsonify({'error': f'Undo operation failed: {str(e)}'}), 500

@app.route('/save-text', methods=['POST'])
def save_text():
    """Endpoint to save selected text to Google Drive or locally as a .docx file."""
    data = request.json
    text = data.get('text')
    file_name = data.get('file_name')
    print("file_name", file_name)

    save_location = data.get('save_location', 'local')  # Default to local
    if not text:
        return jsonify({'error': 'No text provided'}), 400
    
    try:
        if save_location.lower() == 'google':
            # Authenticate with Google Drive
            service = authenticate_google_drive()
            
            # Create a Google Doc with the selected text

            file_id = create_google_doc(service, text, file_name)
            
            return jsonify({
                'message': 'Text saved to Google Drive',
                'file_id': file_id
            }), 200
        
        elif save_location.lower() == 'local':
            # Create a local .docx file
            file_path = create_local_docx(text, file_name = file_name, save_location = save_location)
            
            return jsonify({
                'message': 'Text saved locally as a .docx file',
                'file_path': file_path
            }), 200
        
        else:
            return jsonify({'error': 'Invalid save location. Use "google" or "local".'}), 400
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)